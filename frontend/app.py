"""
실손의료보험 약관 초안 작성 에이전트 — Streamlit UI
백엔드: http://localhost:8000 (FastAPI + LangGraph)

구조:
  STEP 1. 기본 설정     (보험사 선택, 상품명)
  STEP 2. 상품 설계 조건 (보험기간, 납입, 갱신, 가입나이)
  STEP 3. 보장 조건     (보장 종목, 비급여, 한도, 자기부담)
  STEP 4. 확인 및 생성  (입력값 요약 + 생성 버튼)
  결과   : 약관 / 상품설명서 / 사업방법서 3탭 항상 표시
"""

import io
import json
import os
import re
from uuid import uuid4

import requests
import sseclient
from dotenv import load_dotenv

load_dotenv()

import streamlit as st

# ────────────────────────────────────────────────────────────────────────────
# 상수
# ────────────────────────────────────────────────────────────────────────────

MAX_ITERATIONS = 3
BACKEND_URL = os.getenv("BACKEND_URL", "http://localhost:8000")

STEPS_GENERATE = [
    "기본 설정",
    "상품 설계 조건",
    "보장 조건",
    "확인 및 생성",
]
STEPS_REVISE = [
    "기본 설정",
    "약관 업로드",
    "보장 조건 (선택)",
    "확인 및 검증",
]
WORKFLOW_MODE_GENERATE = "AI 조건 기반 생성"
WORKFLOW_MODE_REVISE   = "직접 작성한 약관 검증·수정"

# ────────────────────────────────────────────────────────────────────────────
# 3사 기본값 (JSON 데이터 기반)
# ────────────────────────────────────────────────────────────────────────────

COMPANY_DEFAULTS = {
    "삼성화재": {
        "product_name":             "무배당 삼성화재 다이렉트 실손의료비보험(2605.1)",
        "product_version":          "2605.1",
        "join_age_min":             "0세",
        "join_age_max":             "65세",
        "fetal_enrollment":         "가능",
        "max_coverage_age":         100,
        "policy_period":            "1년 만기",
        "premium_payment_period":   "전기납",
        "premium_payment_cycle":    "월납",
        "renewal_type":             "갱신형",
        "renewal_period":           "1년",
        "max_renewal_count":        4,
        "reinstatement_cycle":      "5년",
        "policy_loan":              "가능",
        "coverage_limit_basic":     "5천만원",
        "coverage_limit_noncovered":"5천만원",
        "coverage_limit_dosu":      "350만원",
        "coverage_limit_injection": "250만원",
        "coverage_limit_mri":       "300만원",
        "outpatient_limit":         "20만원",
        "deductible_hospital":      "1만원 또는 보장대상의료비의 20% 중 큰 금액",
        "deductible_major":         "2만원 또는 보장대상의료비의 20% 중 큰 금액",
    },
    "현대해상": {
        "product_name":             "무배당 현대해상다이렉트실손의료비보장보험(갱신형)(Hi2605)",
        "product_version":          "Hi2605",
        "join_age_min":             "태아",
        "join_age_max":             "60세",
        "fetal_enrollment":         "가능",
        "max_coverage_age":         100,
        "policy_period":            "1년 만기",
        "premium_payment_period":   "전기납",
        "premium_payment_cycle":    "월납",
        "renewal_type":             "갱신형",
        "renewal_period":           "1년",
        "max_renewal_count":        4,
        "reinstatement_cycle":      "5년",
        "policy_loan":              "가능",
        "coverage_limit_basic":     "5천만원",
        "coverage_limit_noncovered":"5천만원",
        "coverage_limit_dosu":      "350만원",
        "coverage_limit_injection": "250만원",
        "coverage_limit_mri":       "300만원",
        "outpatient_limit":         "20만원",
        "deductible_hospital":      "1만원 또는 보장대상의료비의 20% 중 큰 금액",
        "deductible_major":         "2만원 또는 보장대상의료비의 20% 중 큰 금액",
    },
    "DB손해보험": {
        "product_name":             "무배당프로미라이프실손의료비보험",
        "product_version":          "2605",
        "join_age_min":             "5세",
        "join_age_max":             "99세",
        "fetal_enrollment":         "가능",
        "max_coverage_age":         100,
        "policy_period":            "1년 만기",
        "premium_payment_period":   "전기납",
        "premium_payment_cycle":    "월납",
        "renewal_type":             "갱신형",
        "renewal_period":           "1년",
        "max_renewal_count":        4,
        "reinstatement_cycle":      "5년",
        "policy_loan":              "가능",
        "coverage_limit_basic":     "5천만원",
        "coverage_limit_noncovered":"5천만원",
        "coverage_limit_dosu":      "350만원",
        "coverage_limit_injection": "250만원",
        "coverage_limit_mri":       "200만원",
        "outpatient_limit":         "20만원",
        "deductible_hospital":      "1만원 또는 보장대상의료비의 20% 중 큰 금액",
        "deductible_major":         "2만원 또는 보장대상의료비의 20% 중 큰 금액",
    },
}

COVERAGE_LIMIT_OPTIONS  = ["3천만원", "5천만원"]
DOSU_LIMIT_OPTIONS      = ["250만원", "350만원", "500만원"]
INJECTION_LIMIT_OPTIONS = ["150만원", "250만원", "350만원"]
MRI_LIMIT_OPTIONS       = ["150만원", "200만원", "300만원"]
OUTPATIENT_LIMIT_OPTIONS= ["10만원",  "20만원",  "30만원"]

DEDUCTIBLE_OPTIONS = [
    "1만원 또는 보장대상의료비의 20% 중 큰 금액",
    "2만원 또는 보장대상의료비의 20% 중 큰 금액",
    "15% 정률제",
    "20% 정률제",
]

NODE_EMOJI = {
    "coordinator": "🔍",
    "planner":     "📋",
    "supervisor":  "🎯",
    "generation":  "✍️",
    "compliance":  "⚖️",
    "edit":        "✏️",
    "revise":      "🔧",
}

REGULATORY_RISK_MODEL = "solar-pro"
SAMPLE_CLARITY_DOCUMENTS = {
    "약관": (
        "제5조 보험금 지급\n"
        "회사가 필요하다고 인정하는 경우 보험금을 지급합니다.\n"
        "보험금 지급 사유에 해당하고 필요한 서류 제출 및 심사 절차를 거친 경우 보험금을 지급합니다.\n\n"
        "제6조 보장 한도 및 자기부담금\n"
        "보험금은 연간 보장 한도 및 자기 부담금 공제 후 지급합니다.\n"
        "약관에서 정한 면책 사유 또는 이에 준하는 사유에 해당하는 경우 보험금을 지급하지 않습니다."
    ),
    "상품설명서": (
        "보장 내용\n"
        "이 상품은 다양한 의료비를 폭넓게 보장받을 수 있는 상품입니다.\n"
        "주요 의료비에 대해 보장받을 수 있으나 일부 항목은 보장하지 않을 수 있습니다.\n\n"
        "보험료 및 자기부담금\n"
        "보험료 부담을 줄일 수 있으며 일정 비용만 부담하면 보장이 가능합니다.\n"
        "갱신 시 보험료가 달라질 수 있습니다.\n\n"
        "해지 및 환급금\n"
        "계약을 중도해지하는 경우 환급금은 달라질 수 있습니다.\n"
        "가입 전 주요 유의사항을 확인하시기 바랍니다."
    ),
    "사업방법서": (
        "계약 인수\n"
        "회사가 필요하다고 판단하는 경우 가입이 제한될 수 있습니다.\n"
        "심사 결과에 따라 조건부 인수가 적용될 수 있습니다.\n\n"
        "보험금 지급 심사\n"
        "추가 확인이 필요한 경우 보험금 지급이 제한될 수 있습니다.\n"
        "회사의 내부 기준에 따라 지급 여부를 결정할 수 있습니다.\n\n"
        "운영 기준\n"
        "예외적으로 달리 적용할 수 있으며 필요 시 기준을 변경할 수 있습니다."
    ),
}

# ────────────────────────────────────────────────────────────────────────────
# 페이지 설정
# ────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="실손의료보험 초안 작성 에이전트",
    page_icon="📄",
    layout="wide",
)

# ────────────────────────────────────────────────────────────────────────────
# CSS
# ────────────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
.step-bar {
    display: flex;
    align-items: center;
    gap: 0;
    margin-bottom: 1.4rem;
}
.step-item {
    display: flex;
    align-items: center;
    gap: 0.4rem;
    font-size: 0.82rem;
    font-weight: 600;
    color: #94a3b8;
    white-space: nowrap;
}
.step-item.active { color: #2563eb; }
.step-item.done   { color: #16a34a; }
.step-circle {
    width: 26px;
    height: 26px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.75rem;
    font-weight: 700;
    background: #e2e8f0;
    color: #64748b;
    flex-shrink: 0;
}
.step-item.active .step-circle { background: #2563eb; color: white; }
.step-item.done   .step-circle { background: #16a34a; color: white; }
.step-connector {
    flex: 1;
    height: 2px;
    background: #e2e8f0;
    margin: 0 6px;
    min-width: 16px;
}
.step-connector.done { background: #16a34a; }
.summary-card {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    margin-bottom: 0.7rem;
}
.summary-card h4 {
    margin: 0 0 0.6rem 0;
    font-size: 0.85rem;
    color: #64748b;
    text-transform: uppercase;
    letter-spacing: 0.05em;
}
.summary-row {
    display: flex;
    justify-content: space-between;
    font-size: 0.88rem;
    padding: 0.18rem 0;
    border-bottom: 1px solid #f1f5f9;
}
.summary-row:last-child { border-bottom: none; }
.summary-key { color: #64748b; }
.summary-val { font-weight: 600; color: #1e293b; }
</style>
""", unsafe_allow_html=True)


# ────────────────────────────────────────────────────────────────────────────
# 세션 초기화
# ────────────────────────────────────────────────────────────────────────────

# Streamlit은 한 run에서 렌더링되지 않은 위젯의 session_state 값을 다음 run에서
# 삭제한다. 이 앱은 STEP마다 다른 위젯만 렌더링하는 마법사 UI이므로, 아래 키들은
# 사용자가 다른 STEP으로 이동하는 순간 사라지고 _init_session()이 회사 기본값으로
# 되돌려버린다. _snapshot_form_fields()로 매 run마다 백업해 두고, _init_session()이
# 키가 사라졌을 때 하드코딩된 기본값이 아니라 이 백업에서 복원하도록 한다.
_PERSIST_FIELDS = [
    "insurance_company", "insurance_type", "product_name", "product_version",
    "premium_payment_cycle", "max_renewal_count", "join_age_min", "join_age_max",
    "max_coverage_age", "fetal_enrollment", "policy_loan",
    "noncovered_enabled", "three_major_enabled",
    "coverage_limit_basic", "coverage_limit_noncovered", "coverage_limit_dosu",
    "coverage_limit_injection", "coverage_limit_mri", "outpatient_limit",
    "deductible_hospital", "deductible_major",
]


def _snapshot_form_fields():
    """현재 session_state에 살아있는 폼 필드 값을 영구 백업에 복사.
    매 run의 스텝 렌더링 직후 호출해야 한다."""
    snap = st.session_state.setdefault("_form_snapshot", {})
    for k in _PERSIST_FIELDS:
        if k in st.session_state:
            snap[k] = st.session_state[k]


def _field(key: str, default=None):
    """session_state에 키가 없으면(다른 STEP으로 이동해 위젯이 GC된 경우)
    폼 스냅샷에서 복원해 반환한다. build_request 등 최종 읽기 지점에서 사용."""
    if key in st.session_state:
        return st.session_state[key]
    return st.session_state.get("_form_snapshot", {}).get(key, default)


# ────────────────────────────────────────────────────────────────────────────
# 영속 위젯 헬퍼
#
# 위 _PERSIST_FIELDS 키를 위젯의 key=로 직접 사용하면 안 된다 — 다른 STEP으로
# 이동해 위젯이 다시 렌더링되지 않으면 Streamlit이 다음 run에서 그 값을 비우거나
# 이전에 그 위젯이 가지고 있던 값으로 되돌리는 사례가 실측으로 확인됐다.
# 따라서 위젯은 매번 별도의 내부 key(_w_<field>)로 새로 만들고, 항상 canonical
# 값(_field)으로 명시적으로 시드한 뒤, 반환값을 canonical 키에 즉시 되돌려 쓴다.
# ────────────────────────────────────────────────────────────────────────────

def _p_text_input(label, field, default="", **kwargs):
    wkey = f"_w_{field}"
    if wkey not in st.session_state:
        st.session_state[wkey] = _field(field, default)
    val = st.text_input(label, key=wkey, **kwargs)
    st.session_state[field] = val
    return val


def _p_number_input(label, field, default=0, **kwargs):
    wkey = f"_w_{field}"
    if wkey not in st.session_state:
        st.session_state[wkey] = _field(field, default)
    val = st.number_input(label, key=wkey, **kwargs)
    st.session_state[field] = val
    return val


def _p_selectbox(label, options, field, default=None, **kwargs):
    wkey = f"_w_{field}"
    if wkey not in st.session_state:
        cur = _field(field, default if default is not None else options[0])
        st.session_state[wkey] = cur if cur in options else options[0]
    val = st.selectbox(label, options, key=wkey, **kwargs)
    st.session_state[field] = val
    return val


def _p_radio(label, options, field, default=None, **kwargs):
    wkey = f"_w_{field}"
    if wkey not in st.session_state:
        cur = _field(field, default if default is not None else options[0])
        st.session_state[wkey] = cur if cur in options else options[0]
    val = st.radio(label, options, key=wkey, **kwargs)
    st.session_state[field] = val
    return val


def _p_checkbox(label, field, default=False, **kwargs):
    wkey = f"_w_{field}"
    if wkey not in st.session_state:
        st.session_state[wkey] = _field(field, default)
    val = st.checkbox(label, key=wkey, **kwargs)
    st.session_state[field] = val
    return val


def _init_session():
    if "current_step" not in st.session_state:
        st.session_state.current_step = 0
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid4())
    if "workflow_mode" not in st.session_state:
        st.session_state.workflow_mode = WORKFLOW_MODE_GENERATE

    if "generated_result" not in st.session_state:
        st.session_state.generated_result = None

    if "regulatory_risk_result" not in st.session_state:
        st.session_state.regulatory_risk_result = None

    snap = st.session_state.get("_form_snapshot", {})
    defaults = COMPANY_DEFAULTS["삼성화재"]
    init_keys = {
        "insurance_company":  "삼성화재",
        "insurance_type":     "기본형 실손의료비보험",
        "policy_period":      defaults["policy_period"],
        "premium_payment_period": defaults["premium_payment_period"],
        "premium_payment_cycle":  defaults["premium_payment_cycle"],
        "renewal_type":       defaults["renewal_type"],
        "renewal_period":     defaults["renewal_period"],
        "max_renewal_count":  defaults["max_renewal_count"],
        "reinstatement_cycle":defaults["reinstatement_cycle"],
        "join_age_min":       defaults["join_age_min"],
        "join_age_max":       defaults["join_age_max"],
        "max_coverage_age":   defaults["max_coverage_age"],
        "fetal_enrollment":   defaults["fetal_enrollment"],
        "policy_loan":        defaults["policy_loan"],
        "product_name":       defaults["product_name"],
        "product_version":    defaults["product_version"],
        "basic_coverage_items":         ["상해급여", "질병급여"],
        "noncovered_enabled":           False,
        "three_major_enabled":          False,
        "noncovered_rider_items":       [],
        "three_major_noncovered_items": [],
        "coverage_limit_basic":         defaults["coverage_limit_basic"],
        "coverage_limit_noncovered":    defaults["coverage_limit_noncovered"],
        "coverage_limit_dosu":          defaults["coverage_limit_dosu"],
        "coverage_limit_injection":     defaults["coverage_limit_injection"],
        "coverage_limit_mri":           defaults["coverage_limit_mri"],
        "outpatient_limit":             defaults["outpatient_limit"],
        "deductible_hospital":          defaults["deductible_hospital"],
        "deductible_major":             defaults["deductible_major"],
    }
    for k, v in init_keys.items():
        if k not in st.session_state:
            st.session_state[k] = snap.get(k, v)


def apply_company_defaults(company: str):
    """회사 변경 시 호출. canonical 키와 위젯 키(_w_<field>)를 모두 갱신해야 한다 —
    위젯이 이미 존재하는 상태에서는 value=/index= 인자가 무시되고 위젯 자신의
    session_state 값이 우선하므로, 위젯 키를 직접 덮어써야 화면에 반영된다."""
    d = COMPANY_DEFAULTS.get(company, COMPANY_DEFAULTS["삼성화재"])
    for f, v in d.items():
        st.session_state[f] = v
        st.session_state[f"_w_{f}"] = v


_init_session()


# ────────────────────────────────────────────────────────────────────────────
# 네비게이션
# ────────────────────────────────────────────────────────────────────────────

def go_prev(): st.session_state.current_step -= 1
def go_next(): st.session_state.current_step += 1


# ────────────────────────────────────────────────────────────────────────────
# 스텝 인디케이터
# ────────────────────────────────────────────────────────────────────────────

def render_step_bar(current: int, steps: list):
    html = '<div class="step-bar">'
    for i, label in enumerate(steps):
        if i < current:
            cls, circle = "done", "✓"
        elif i == current:
            cls, circle = "active", str(i + 1)
        else:
            cls, circle = "", str(i + 1)
        html += f'<div class="step-item {cls}"><div class="step-circle">{circle}</div><span>{label}</span></div>'
        if i < len(steps) - 1:
            conn_cls = "done" if i < current else ""
            html += f'<div class="step-connector {conn_cls}"></div>'
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)


# ────────────────────────────────────────────────────────────────────────────
# 스텝 렌더링
# ────────────────────────────────────────────────────────────────────────────

def render_step1():
    """STEP 1: 기본 설정"""
    st.subheader("🏢 기본 설정")
    st.caption("보험사를 선택하면 해당 보험사의 실제 상품 기본값이 자동으로 채워집니다.")

    _p_selectbox(
        "보험사",
        list(COMPANY_DEFAULTS.keys()),
        "insurance_company",
        on_change=lambda: apply_company_defaults(st.session_state["_w_insurance_company"]),
    )

    st.divider()
    if st.session_state.workflow_mode == WORKFLOW_MODE_GENERATE:
        st.caption("ℹ️ 약관·상품설명서·사업방법서 3개 문서를 한 번에 생성합니다.")
    else:
        st.caption("ℹ️ 업로드한 약관을 검증→수정 반복 후 최종 약관만 출력합니다.")

    st.divider()
    _p_text_input("상품명",          "product_name",    help="보험사 변경 시 자동 업데이트")
    _p_text_input("버전 / 상품코드", "product_version")
    _p_selectbox("보험 유형", ["기본형 실손의료비보험", "특약 포함 실손의료비보험"], "insurance_type")

    return True  # STEP 1은 항상 다음으로 진행 가능


def _extract_text_from_upload(uploaded_file) -> str:
    """업로드 파일(.docx/.txt)에서 텍스트를 추출."""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    return data.decode("utf-8", errors="ignore")


def render_step2_revise():
    """STEP 2 (검증·수정 모드): 사용자 작성 약관 업로드"""
    st.subheader("📤 약관 업로드")
    st.caption(
        "직접 작성한 약관 전문을 .docx 또는 .txt 파일로 업로드하세요. "
        "AI 생성 단계 없이 법규 검증 → 수정을 반복하여 최종 약관을 출력합니다."
    )

    uploaded = st.file_uploader("약관 파일", type=["docx", "txt"], key="_user_doc_upload")
    if uploaded is not None:
        try:
            text = _extract_text_from_upload(uploaded)
            st.session_state.user_document_text = text
            st.session_state.user_document_name = uploaded.name
        except Exception as e:
            st.error(f"파일을 읽는 중 오류가 발생했습니다: {e}")

    text = st.session_state.get("user_document_text", "")
    if text:
        st.success(f"✅ {st.session_state.get('user_document_name', '')} ({len(text):,}자) 업로드 완료")
        with st.expander("📄 업로드된 약관 미리보기"):
            st.text(text[:3000] + ("..." if len(text) > 3000 else ""))
    else:
        st.warning("⚠️ 약관 파일을 업로드해야 다음 단계로 진행할 수 있습니다.")

    return bool(text)


def render_step2():
    """STEP 2: 상품 설계 조건"""
    st.subheader("⚙️ 상품 설계 조건")
    st.caption("사업방법서 기준 항목입니다.")

    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("**📅 보험기간 / 납입**")
        st.text_input("보험기간", value="1년 만기", disabled=True, key="_fp_period")
        st.caption("✔ 실손의료보험 1년 만기 고정 (법령)")

        st.text_input("보험료 납입기간", value="전기납", disabled=True, key="_fp_pay")
        st.caption("✔ 전기납 고정")

        _p_selectbox("보험료 납입주기", ["월납", "3개월납", "6개월납", "연납"], "premium_payment_cycle")

        st.divider()
        st.markdown("**🔁 갱신**")
        st.text_input("갱신 유형", value="갱신형", disabled=True, key="_fp_renewal")
        st.caption("✔ 갱신형 고정 (현행 제도)")

        st.text_input("갱신 주기", value="1년", disabled=True, key="_fp_rcycle")
        st.caption("✔ 1년 고정")

        _p_number_input("최대 갱신 횟수", "max_renewal_count", default=4, min_value=1, max_value=10, step=1)

        st.text_input("재가입 주기", value="5년", disabled=True, key="_fp_reinstate")
        st.caption("✔ 5년 고정")

    with col_b:
        st.markdown("**👤 가입 조건**")
        company = _field("insurance_company", "삼성화재")
        age_hints = {"삼성화재": "0세 / 65세", "현대해상": "태아 / 60세", "DB손해보험": "5세 / 99세"}
        st.caption(f"ℹ️ {company} 기본: {age_hints.get(company, '')}")

        defaults = COMPANY_DEFAULTS.get(company, COMPANY_DEFAULTS["삼성화재"])
        _p_text_input(
            "가입 최소 나이",
            "join_age_min",
            default=defaults["join_age_min"],
            placeholder=defaults["join_age_min"],
        )
        _p_text_input(
            "가입 최대 나이",
            "join_age_max",
            default=defaults["join_age_max"],
            placeholder=defaults["join_age_max"],
        )
        _p_number_input("최대 보장 나이 (세)", "max_coverage_age", default=100, min_value=70, max_value=120, step=1)

        st.divider()
        st.markdown("**🍼 특수 조건**")
        _p_radio("태아 가입 가능 여부", ["가능", "불가"], "fetal_enrollment", default="가능", horizontal=True)
        _p_radio("보험계약대출 가능 여부", ["가능", "불가"], "policy_loan", default="가능", horizontal=True)

    # 필수 항목 유효성 검사
    missing = [
        label for label, key in [("가입 최소 나이", "join_age_min"), ("가입 최대 나이", "join_age_max")]
        if not st.session_state.get(key, "").strip()
    ]
    if missing:
        st.error(f"⚠️ 필수 항목을 입력해주세요: {', '.join(missing)}")
        return False
    return True


def render_step3(required: bool = True):
    """STEP 3: 보장 조건

    required=False (검증·수정 모드): 입력하지 않아도 다음 단계로 진행 가능.
    미입력 시 compliance 검증은 보편적 법규 기준만으로 수행된다.
    """
    st.subheader("🏥 보장 조건")
    if not required:
        st.caption("ℹ️ 선택 입력입니다. 비워두면 보편적 법규 기준으로만 검증합니다.")

    # 기본 보장 종목
    label = "**기본 보장 종목** (약관 제1조 기준, 1개 이상 필수)" if required else "**기본 보장 종목** (선택)"
    st.markdown(label)
    basic_items = st.session_state.get("basic_coverage_items", [])
    col_b1, col_b2 = st.columns(2)
    with col_b1:
        injury  = st.checkbox("상해급여 실손의료비", value="상해급여" in basic_items, key="_basic_injury",
                              help="상해로 인한 입원·통원 급여의료비 보상")
    with col_b2:
        disease = st.checkbox("질병급여 실손의료비", value="질병급여" in basic_items, key="_basic_disease",
                              help="질병으로 인한 입원·통원 급여의료비 보상")

    new_basic = (["상해급여"] if injury else []) + (["질병급여"] if disease else [])
    st.session_state.basic_coverage_items = new_basic
    if not new_basic:
        if required:
            st.error("⚠️ 기본 보장 종목을 최소 1개 이상 선택해야 합니다.")
        else:
            st.caption("※ 미선택 시 보장 한도·자기부담금 검증은 생략됩니다.")

    st.divider()

    # 비급여 특약
    st.markdown("**비급여 특약** (선택사항)")
    noncov = _p_checkbox("비급여 실손의료비 특약 포함", "noncovered_enabled",
                         help="중증·비중증 비급여 의료비 보상 특약")
    if noncov:
        st.session_state.noncovered_rider_items = [
            "[갱신형]중증 비급여 실손의료비 특별약관(상해)",
            "[갱신형]중증 비급여 실손의료비 특별약관(질병)",
            "[갱신형]비중증 비급여 실손의료비 특별약관(상해)",
            "[갱신형]비중증 비급여 실손의료비 특별약관(질병)",
        ]
        three_major = _p_checkbox("3대 비급여 세부항목 포함 (도수치료·주사료·MRI)", "three_major_enabled")
        if three_major:
            st.session_state.three_major_noncovered_items = [
                "[갱신형]중증 비급여 실손의료비 특별약관(3대비급여)",
                "[갱신형]비중증 비급여 자기공명영상진단보장 특별약관",
            ]
        else:
            st.session_state.three_major_noncovered_items = []
    else:
        st.session_state.noncovered_rider_items       = []
        st.session_state.three_major_noncovered_items = []
        st.session_state.three_major_enabled          = False

    st.divider()

    # 보험가입금액 한도
    st.markdown("**보험가입금액 한도**")
    col_l1, col_l2 = st.columns(2)
    with col_l1:
        _sel(st, "급여 연간 한도",  COVERAGE_LIMIT_OPTIONS,   "coverage_limit_basic")
        _sel(st, "통원 1회 한도",   OUTPATIENT_LIMIT_OPTIONS, "outpatient_limit",
             help="약관 제6조 ⑤항")
    with col_l2:
        _sel(st, "비급여 연간 한도", COVERAGE_LIMIT_OPTIONS, "coverage_limit_noncovered",
             disabled=not noncov,
             help="비급여 특약 선택 시 활성화" if not noncov else "")
        if not noncov:
            st.caption("※ 비급여 특약 선택 시 활성화됩니다.")

    if noncov and st.session_state.get("three_major_enabled", False):
        st.markdown("**3대 비급여 세부 한도**")
        col_3a, col_3b, col_3c = st.columns(3)
        with col_3a: _sel(st, "도수치료 한도", DOSU_LIMIT_OPTIONS,      "coverage_limit_dosu")
        with col_3b: _sel(st, "주사료 한도",   INJECTION_LIMIT_OPTIONS, "coverage_limit_injection")
        with col_3c:
            _sel(st, "MRI/MRA 한도", MRI_LIMIT_OPTIONS, "coverage_limit_mri",
                 help="DB손해보험 200만원 / 삼성·현대 300만원")

    st.divider()

    # 자기부담금
    st.markdown("**자기부담금 (공제금액) 규칙** — 약관 제3조 〈표1〉")
    col_d1, col_d2 = st.columns(2)
    with col_d1: _sel(st, "일반 의료기관 통원", DEDUCTIBLE_OPTIONS, "deductible_hospital")
    with col_d2: _sel(st, "상급종합병원 통원",  DEDUCTIBLE_OPTIONS, "deductible_major",
                      index_default=1)

    return bool(new_basic) if required else True


def _sel(parent, label, options, key, disabled=False, help="", index_default=0):
    """selectbox 헬퍼 — _p_selectbox 기반으로 STEP 이동 후에도 값을 안전하게 유지"""
    default = options[index_default] if 0 <= index_default < len(options) else options[0]
    _p_selectbox(label, options, key, default=default, disabled=disabled, help=help)


def render_step4():
    """STEP 4: 확인 및 생성/검증"""
    s = st.session_state
    is_revise_mode = s.get("workflow_mode") == WORKFLOW_MODE_REVISE

    def card(title, rows):
        html = f'<div class="summary-card"><h4>{title}</h4>'
        for k, v in rows:
            html += (f'<div class="summary-row">'
                     f'<span class="summary-key">{k}</span>'
                     f'<span class="summary-val">{v}</span></div>')
        html += "</div>"
        st.markdown(html, unsafe_allow_html=True)

    if is_revise_mode:
        st.subheader("✅ 확인 및 검증")
        st.caption("업로드한 약관을 법규 검증→수정 반복 후 최종 약관만 출력합니다.")

        card("기본 설정", [
            ("보험사",    _field("insurance_company", "-")),
            ("상품명",    _field("product_name", "-")),
            ("버전/코드", _field("product_version", "-")),
        ])
        card("업로드된 약관", [
            ("파일명", s.get("user_document_name", "-")),
            ("글자 수", f"{len(s.get('user_document_text', '')):,}자"),
        ])
        basic = s.get("basic_coverage_items", [])
        card("보장 조건 (선택 입력)", [
            ("기본 보장 종목", ", ".join(basic) if basic else "미입력 — 보편적 법규만 검증"),
        ])
        return

    st.subheader("✅ 입력 확인")
    st.caption("아래 내용으로 초안을 생성합니다. 수정이 필요하면 이전 단계로 돌아가세요.")

    card("기본 설정", [
        ("보험사",    _field("insurance_company", "-")),
        ("상품명",    _field("product_name", "-")),
        ("버전/코드", _field("product_version", "-")),
        ("보험 유형", _field("insurance_type", "-")),
    ])

    card("상품 설계 조건", [
        ("보험기간",       s.get("policy_period", "-")),
        ("납입기간/주기",  f"{s.get('premium_payment_period','-')} / {_field('premium_payment_cycle','-')}"),
        ("갱신",           f"{s.get('renewal_type','-')} ({s.get('renewal_period','-')}, 최대 {_field('max_renewal_count','-')}회)"),
        ("재가입 주기",    s.get("reinstatement_cycle", "-")),
        ("가입나이",       f"{_field('join_age_min','-')} ~ {_field('join_age_max','-')}"),
        ("최대 보장나이",  f"{_field('max_coverage_age','-')}세"),
        ("태아 가입",      _field("fetal_enrollment", "-")),
        ("보험계약대출",   _field("policy_loan", "-")),
    ])

    basic        = s.get("basic_coverage_items", [])
    noncov_items = s.get("noncovered_rider_items", [])
    three_items  = s.get("three_major_noncovered_items", [])
    cov_rows = [
        ("기본 보장 종목",  ", ".join(basic) if basic else "미선택"),
        ("비급여 특약",     "포함" if noncov_items else "미포함"),
        ("3대 비급여 특약", "포함" if three_items  else "미포함"),
        ("급여 연간 한도",  _field("coverage_limit_basic",  "-")),
        ("통원 1회 한도",   _field("outpatient_limit",       "-")),
    ]
    if noncov_items:
        cov_rows.append(("비급여 연간 한도", _field("coverage_limit_noncovered", "-")))
    if three_items:
        cov_rows += [
            ("도수치료 한도", _field("coverage_limit_dosu",      "-")),
            ("주사료 한도",   _field("coverage_limit_injection",  "-")),
            ("MRI/MRA 한도",  _field("coverage_limit_mri",        "-")),
        ]
    cov_rows += [
        ("일반 통원 공제", _field("deductible_hospital", "-")),
        ("종합병원 공제",  _field("deductible_major",    "-")),
    ]
    card("보장 조건", cov_rows)


# ────────────────────────────────────────────────────────────────────────────
# 백엔드 요청 빌더
# ────────────────────────────────────────────────────────────────────────────

def build_request(model=None) -> dict:
    s = st.session_state
    return {
        "document_request": {
            "document_type":     "전체",   # 약관·상품설명서·사업방법서 일괄 생성
            "insurance_company": _field("insurance_company", "삼성화재"),
            "insurance_type":    _field("insurance_type", "기본형 실손의료비보험"),
            "product_name":      _field("product_name", ""),
            "product_version":   _field("product_version", ""),
            "dividend_type":     "무배당",
        },
        "product_design_conditions": {
            "policy_period":          s.get("policy_period", "1년 만기"),
            "premium_payment_period": s.get("premium_payment_period", "전기납"),
            "premium_payment_cycle":  _field("premium_payment_cycle", "월납"),
            "renewal_type":           s.get("renewal_type", "갱신형"),
            "renewal_period":         s.get("renewal_period", "1년"),
            "max_renewal_count":      _field("max_renewal_count", 4),
            "reinstatement_cycle":    s.get("reinstatement_cycle", "5년"),
            "max_coverage_age":       _field("max_coverage_age", 100),
            "join_age_min":           _field("join_age_min", "0세"),
            "join_age_max":           _field("join_age_max", "65세"),
            "fetal_enrollment":       _field("fetal_enrollment", "가능"),
            "policy_loan":            _field("policy_loan", "가능"),
        },
        "coverage_conditions": {
            "basic_coverage_items":         s.get("basic_coverage_items", []),
            "noncovered_rider_items":       s.get("noncovered_rider_items", []),
            "three_major_noncovered_items": s.get("three_major_noncovered_items", []),
            "coverage_limit": {
                "급여":     _field("coverage_limit_basic",       "5천만원"),
                "비급여":   _field("coverage_limit_noncovered",  "5천만원"),
                "도수치료": _field("coverage_limit_dosu",        "350만원"),
                "주사료":   _field("coverage_limit_injection",   "250만원"),
                "MRI":      _field("coverage_limit_mri",         "300만원"),
            },
            "outpatient_limit": _field("outpatient_limit", "20만원"),
            "deductible_rule": {
                "일반_의료기관": _field("deductible_hospital", ""),
                "상급종합병원":  _field("deductible_major",    ""),
            },
        },
        "session_id": s.session_id,
        "model": model,
    }


def build_revise_request(model=None) -> dict:
    """검증·수정 모드 요청 빌더 — 사용자 작성 약관 전문을 user_document로 전달.
    보장 조건은 STEP 3에서 입력한 만큼만 채우고, 미입력 시 빈 값으로 보내
    compliance가 보편적 법규 기준으로만 검증하도록 한다."""
    s = st.session_state
    basic_items = s.get("basic_coverage_items", [])
    return {
        "document_request": {
            "document_type":     "약관",
            "insurance_company": _field("insurance_company", ""),
            "insurance_type":    _field("insurance_type", "기본형 실손의료비보험"),
            "product_name":      _field("product_name", ""),
            "product_version":   _field("product_version", ""),
            "dividend_type":     "무배당",
        },
        "product_design_conditions": {},
        "coverage_conditions": {
            "basic_coverage_items":         basic_items,
            "noncovered_rider_items":       s.get("noncovered_rider_items", []),
            "three_major_noncovered_items": s.get("three_major_noncovered_items", []),
            "coverage_limit": {
                "급여":     _field("coverage_limit_basic", ""),
                "비급여":   _field("coverage_limit_noncovered", ""),
                "도수치료": _field("coverage_limit_dosu", ""),
                "주사료":   _field("coverage_limit_injection", ""),
                "MRI":      _field("coverage_limit_mri", ""),
            } if basic_items else {},
            "outpatient_limit": _field("outpatient_limit", ""),
            "deductible_rule": {
                "일반_의료기관": _field("deductible_hospital", ""),
                "상급종합병원":  _field("deductible_major", ""),
            } if basic_items else {},
        },
        "user_document": s.get("user_document_text", ""),
        "session_id": s.session_id,
        "model": model,
    }


# ────────────────────────────────────────────────────────────────────────────
# 위반 하이라이트
# ────────────────────────────────────────────────────────────────────────────


async def _run_parallel() -> list[dict]:
    async with aiohttp.ClientSession() as session:
        return list(await asyncio.gather(*[_post_one(session, m) for m in MODELS]))


def extract_text_from_nested_result(obj, max_depth: int = 4) -> tuple[str, str]:
    best_text = ""
    best_path = ""

    def visit(value, path: str, depth: int) -> None:
        nonlocal best_text, best_path
        if depth > max_depth:
            return
        if isinstance(value, str):
            text = value.strip()
            if len(text) > len(best_text):
                best_text = text
                best_path = path
            return
        if isinstance(value, dict):
            for key, nested in value.items():
                visit(nested, f"{path}.{key}" if path else str(key), depth + 1)
            return
        if isinstance(value, list):
            text_items = [item.strip() for item in value if isinstance(item, str) and item.strip()]
            if text_items:
                combined = "\n".join(text_items)
                if len(combined) > len(best_text):
                    best_text = combined
                    best_path = path
            for index, nested in enumerate(value[:10]):
                visit(nested, f"{path}[{index}]", depth + 1)

    visit(obj, "", 0)
    return best_text, best_path


def extract_generated_document_text(result: dict, document_type: str = "약관") -> tuple[str, str]:
    fields_by_document_type = {
        "약관": ["final_content", "final_document", "draft_content", "content", "generated_draft", "generated_text"],
        "상품설명서": ["product_description", "description", "product_description_content", "final_document"],
        "사업방법서": ["business_method", "business_method_content", "operation_method"],
    }
    candidate_fields = fields_by_document_type.get(document_type, []) + [
        "final_content",
        "final_document",
        "draft_content",
        "content",
        "generated_draft",
        "generated_text",
        "product_description",
        "business_method",
        "answer",
    ]
    for key in candidate_fields:
        value = result.get(key)
        if isinstance(value, str) and len(value.strip()) > 20:
            return value.strip(), key
        if isinstance(value, (dict, list)):
            nested_text, nested_path = extract_text_from_nested_result(value, max_depth=3)
            if len(nested_text.strip()) > 20:
                return nested_text.strip(), f"{key}.{nested_path}" if nested_path else key
    nested_text, nested_path = extract_text_from_nested_result(result)
    if len(nested_text.strip()) > 20:
        return nested_text.strip(), nested_path or "nested"
    return "", ""


def run_regulatory_risk_simulation(
    draft_content: str,
    document_type: str = "약관",
    selected_document_field: str = "",
) -> dict:
    response = requests.post(
        f"{BACKEND_URL}/regulatory-risk-simulation",
        json={
            "document_type": document_type,
            "draft_content": draft_content,
            "model": REGULATORY_RISK_MODEL,
            "max_findings": 5,
            "use_llm": True,
            "use_full_compliance_agent": False,
        },
        timeout=120,
    )
    response.raise_for_status()
    result = response.json()
    if selected_document_field:
        result["_frontend_selected_document_field"] = selected_document_field
        result["_frontend_document_text_length"] = len(draft_content or "")
        result["_frontend_document_text_preview"] = (draft_content or "")[:500]
    return result


def render_regulatory_risk_result(result: dict) -> None:
    findings = result.get("regulatory_risk_simulation_findings", [])
    summary = result.get("regulatory_risk_simulation_summary", {})
    source_label = "Neon DB" if summary.get("gray_zone_type_source") == "db" else "기본 평가 기준"
    llm_label = "Upstage Solar Pro 사용" if summary.get("llm_used") else "기본 문장 기준 사용"
    document_type = summary.get("document_type") or "약관"
    matched_labels = sorted(
        {
            item.get("gray_zone_risk_label")
            for item in findings
            if item.get("gray_zone_risk_label")
        }
    )

    st.info(
        "본 기능은 소비자보호 관점에서 초안 문장의 명확성을 평가하고, 더 명확한 문장 제안을 함께 제공합니다. "
        "본 결과는 법적 적법성 판단이나 Compliance 검증 결과가 아닙니다."
    )
    st.markdown("### 평가 결과 요약")
    col_count, col_doc, col_source, col_types, col_llm = st.columns(5)
    col_count.metric("평가 후보", f"{len(findings)}건")
    col_doc.metric("문서 유형", document_type)
    col_source.metric("평가 기준", source_label)
    col_types.metric("평가 유형", f"{summary.get('gray_zone_type_count', 0)}개")
    col_llm.metric("LLM 보강", "사용" if summary.get("llm_used") else "미사용")
    st.caption(
        f"{llm_label} · 제목형 라인 제외 {summary.get('excluded_heading_count', 0)}건 · "
        "원문 제목이나 조항명만으로는 평가 결과를 만들지 않습니다."
    )
    if matched_labels:
        st.markdown(f"**적용된 평가 항목:** {', '.join(matched_labels)}")

    if not findings:
        st.warning(
            "현재 문서에서는 DB 평가 기준과 직접 매칭되는 불명확 표현 후보가 발견되지 않았습니다.\n\n"
            "이는 문서가 법적으로 적합하다는 판단이 아니라, 현재 등록된 평가 유형과 직접 매칭되는 표현이 없다는 의미입니다."
        )
        st.info(
            "시연용으로 문서유형별 예시 문서를 실행하면 평가 결과가 어떤 방식으로 표시되는지 확인할 수 있습니다."
        )
        sample_text = SAMPLE_CLARITY_DOCUMENTS.get(document_type)
        if sample_text and st.button(
            f"현재 문서유형 예시로 평가 실행 ({document_type})",
            key=f"run_zero_sample_{document_type}",
            use_container_width=True,
        ):
            st.session_state.regulatory_risk_result = run_regulatory_risk_simulation(
                sample_text,
                document_type=document_type,
                selected_document_field=f"sample:{document_type}",
            )
            st.rerun()
        with st.expander("개발자용 미매칭 후보 보기", expanded=False):
            st.write("body_candidate_count", summary.get("body_candidate_count", 0))
            st.write("sample_unmatched_body_candidates", summary.get("sample_unmatched_body_candidates", []))

    for index, item in enumerate(findings, 1):
        label = item.get("gray_zone_risk_label") or item.get("gray_zone_risk_type", "")
        unclear_points = item.get("unclear_points_in_source") or item.get("matched_patterns", [])
        if isinstance(unclear_points, list):
            unclear_points_text = ", ".join(str(point) for point in unclear_points if str(point).strip())
        else:
            unclear_points_text = str(unclear_points or "")
        with st.expander(f"평가 결과 {index}. {label}", expanded=index == 1):
            st.markdown(f"**원문 문장**\n\n{item.get('source_text', '')}")
            st.markdown(f"**원문 내 불명확 가능 지점**\n\n{unclear_points_text}")
            st.write("평가 유형", item.get("gray_zone_risk_type", ""))
            st.markdown(
                "**불명확하게 해석될 수 있는 표현 예시**\n\n"
                f"{item.get('ambiguous_expression_example', '')}"
            )
            st.markdown(f"**불명확성 사유**\n\n{item.get('why_ambiguous', '')}")
            st.markdown(f"**소비자 오해 가능성**\n\n{item.get('consumer_confusion_point', '')}")
            st.markdown(f"**명확화 기준**\n\n{item.get('safe_rewrite_guideline', '')}")
            st.markdown(f"**명확한 문장 제안**\n\n{item.get('strengthened_safe_sentence', '')}")
            if item.get("nearest_heading"):
                st.caption(f"가까운 제목: {item.get('nearest_heading')}")
            with st.expander("분류 상세 보기", expanded=False):
                st.write("matched_patterns", item.get("matched_patterns", []))
                st.write("final_classification", item.get("final_classification"))
                st.write("gray_zone_classification", item.get("gray_zone_classification"))
                st.write("classification_basis", item.get("classification_basis"))
                st.write("LLM 보강 상태", item.get("llm_judgement", {}))

    with st.expander("문장 명확성 평가 보고서 미리보기"):
        st.markdown(result.get("regulatory_risk_simulation_report", ""))
    with st.expander("명확한 문장 제안 보고서 미리보기"):
        st.markdown(result.get("safe_alternative_report", ""))
    with st.expander("개발자용 원본 응답 보기", expanded=False):
        st.json(result)


# ── 위반 하이라이트 ───────────────────────────────────────────────────────────
def apply_violation_highlights(content: str, violations: list) -> str:
    for v in violations:
        original = v.get("original_text", "")
        if not original or original not in content:
            continue
        annotation = f"⚠️ [{v['type']}] {v['legal_basis']}: {v['fix']}"
        content = content.replace(
            original,
            f'<span style="color:red;">{original}</span>'
            f'<br><span style="color:red;font-size:12px;">{annotation}</span>',
        )
    return content


# ────────────────────────────────────────────────────────────────────────────
# docx 변환
# ────────────────────────────────────────────────────────────────────────────

def _to_docx_bytes(title: str, content: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in content.split("\n"):
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        elif stripped == "":
            doc.add_paragraph()
        else:
            para = doc.add_paragraph()
            for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", stripped)):
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ────────────────────────────────────────────────────────────────────────────
# 결과 패널
# ────────────────────────────────────────────────────────────────────────────

def render_result_panel(result: dict, model_label: str, show_all_docs: bool = True):
    final_status = result.get("status", "")
    action_word = "생성" if show_all_docs else "수정"
    if final_status == "COMPLIANCE_PASSED":
        st.success(
            f"✅ 법규 검토 통과 — {result.get('iteration', '?')}회 완료 "
            f"({model_label} / 세션: {st.session_state.session_id[:8]}…)"
        )
    elif final_status == "MANUAL_REVIEW_REQUIRED":
        st.error(f"⚠️ 최대 {MAX_ITERATIONS}회 재{action_word} 후에도 법규 준수 미달. 수동 검토 필요.")
        if result.get("suggestions"):
            with st.expander("📋 수동 검토 필요 항목"):
                for s in result["suggestions"]:
                    manual = " 🔴 반복 위반" if s.get("requires_manual_review") else ""
                    st.markdown(
                        f"- **[{s['severity']}] {s['type']}**{manual}\n\n"
                        f"  {s['action']}\n\n  > `{s['target_text']}`"
                    )
    elif final_status == "ORCHESTRATOR_ERROR":
        st.error(f"시스템 오류: {result.get('error', '알 수 없는 오류')}")

    if result.get("db_warning"):      st.warning(f"⚠️ {result['db_warning']}")
    if result.get("improvement_note"):st.info(f"📊 {result['improvement_note']}")

    product_name = st.session_state.get("product_name", "보험상품")
    company      = st.session_state.get("insurance_company", "")

    if not show_all_docs:
        # 검증·수정 모드: 약관만 출력 (상품설명서·사업방법서는 생성하지 않음)
        st.markdown("### 📜 최종 약관")
        content = result.get("content", "")
        highlighted = apply_violation_highlights(content, result.get("violations_for_ui", []))
        st.markdown(highlighted, unsafe_allow_html=True)
        accuracy_history = result.get("accuracy_history", [])
        if accuracy_history:
            st.markdown("---")
            col_title, col_download = st.columns([3, 1])
            with col_title:
                st.subheader("📈 법규 준수율 개선 추이")
            with col_download:
                if content:
                    st.download_button(
                        "⬇️ 약관 다운로드 (.docx)",
                        data=_to_docx_bytes(f"{company} {product_name} 약관", content),
                        file_name=f"{company}_{product_name}_약관.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            edit_count = 0
            gen_above_50_count = 0
            for h in accuracy_history:
                is_post_edit = h.get("is_post_edit", False)
                accuracy   = h.get("accuracy", 0)

                if not is_post_edit and accuracy < 50.0:
                    continue

                violations = h.get("violations", 0)
                h_status   = h.get("status", "")
                bar_filled = int(accuracy / 10)
                bar = "█" * bar_filled + "░" * (10 - bar_filled)
                emoji = "🟢" if accuracy >= 80 else ("🟡" if accuracy >= 60 else "🔴")
                if not is_post_edit:
                    gen_above_50_count += 1
                    label = f"🔄 iteration {gen_above_50_count}"
                else:
                    edit_count += 1
                    label = f"✏️ edit {edit_count}회 후 검증"
                st.markdown(
                    f"**{label}**: {emoji} `{bar}` **{accuracy}%** "
                    f"(위반 {violations}건, {h_status})"
                )
        return

    content = result.get("content", "")
    accuracy_history = result.get("accuracy_history", [])
    if accuracy_history:
        st.markdown("---")
        col_title, col_download = st.columns([3, 1])
        with col_title:
            st.subheader("📈 법규 준수율 개선 추이")
        with col_download:
            if content:
                st.download_button(
                    "⬇️ 약관 다운로드 (.docx)",
                    data=_to_docx_bytes(f"{company} {product_name} 약관", content),
                    file_name=f"{company}_{product_name}_약관.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        edit_count = 0
        gen_above_50_count = 0
        for h in accuracy_history:
            is_post_edit = h.get("is_post_edit", False)
            accuracy   = h.get("accuracy", 0)

            if not is_post_edit and accuracy < 50.0:
                continue

            violations = h.get("violations", 0)
            h_status   = h.get("status", "")
            bar_filled = int(accuracy / 10)
            bar = "█" * bar_filled + "░" * (10 - bar_filled)
            emoji = "🟢" if accuracy >= 80 else ("🟡" if accuracy >= 60 else "🔴")
            if not is_post_edit:
                gen_above_50_count += 1
                label = f"🔄 iteration {gen_above_50_count}"
            else:
                edit_count += 1
                label = f"✏️ edit {edit_count}회 후 검증"
            st.markdown(
                f"**{label}**: {emoji} `{bar}` **{accuracy}%** "
                f"(위반 {violations}건, {h_status})"
            )
        st.markdown("---")

    # ── 텍스트 다운로드 버튼 (탭 위) ────────────────────────────────
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("📄 약관 다운로드", result.get("content", ""), file_name="약관.txt", mime="text/plain")
    with col2:
        st.download_button("📄 상품설명서 다운로드", result.get("product_description", ""), file_name="상품설명서.txt", mime="text/plain")
    with col3:
        st.download_button("📄 사업방법서 다운로드", result.get("business_method", ""), file_name="사업방법서.txt", mime="text/plain")

    # ── 명확성 평가 실행 버튼 (탭 위) ───────────────────────────────
    st.caption("생성된 문서에서 소비자가 불명확하게 해석할 수 있는 표현을 점검합니다.")
    clarity_btn = st.button("🔍 명확성 평가 실행", use_container_width=True, key="clarity_eval_btn")
    if clarity_btn:
        docs = {
            "약관": result.get("content", ""),
            "상품설명서": result.get("product_description", ""),
            "사업방법서": result.get("business_method", ""),
        }
        clarity_results = {}
        for doc_type, doc_content in docs.items():
            if doc_content:
                with st.spinner(f"{doc_type} 명확성 평가 중..."):
                    clarity_results[doc_type] = run_regulatory_risk_simulation(
                        draft_content=doc_content,
                        document_type=doc_type,
                        selected_document_field=doc_type,
                    )
        st.session_state.regulatory_risk_result = clarity_results

    # ── 6탭: 문서 3 + 명확성 평가 3 ─────────────────────────────────
    clarity_results = st.session_state.get("regulatory_risk_result", {})
    tab_clause, tab_desc, tab_biz, tab_cl1, tab_cl2, tab_cl3 = st.tabs([
        "📄 약관", "📄 상품설명서", "📄 사업방법서",
        "🔍 명확성:약관", "🔍 명확성:상품설명서", "🔍 명확성:사업방법서",
    ])

    with tab_clause:
        highlighted = apply_violation_highlights(
            result.get("content", ""), result.get("violations_for_ui", [])
        )
        st.markdown(highlighted, unsafe_allow_html=True)

    with tab_desc:
        desc = result.get("product_description", "")
        st.markdown(desc)
        if desc:
            st.download_button(
                "⬇️ 상품설명서 다운로드 (.docx)",
                data=_to_docx_bytes(f"{company} {product_name} 상품설명서", desc),
                file_name=f"{company}_{product_name}_상품설명서.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with tab_biz:
        biz = result.get("business_method", "")
        st.markdown(biz)
        if biz:
            st.download_button(
                "⬇️ 사업방법서 다운로드 (.docx)",
                data=_to_docx_bytes(f"{company} {product_name} 사업방법서", biz),
                file_name=f"{company}_{product_name}_사업방법서.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    for tab, doc_type in zip([tab_cl1, tab_cl2, tab_cl3], ["약관", "상품설명서", "사업방법서"]):
        with tab:
            if clarity_results and doc_type in clarity_results:
                render_regulatory_risk_result(clarity_results[doc_type])
            else:
                st.info("위쪽 '🔍 명확성 평가 실행' 버튼을 클릭하면 평가 결과가 여기에 표시됩니다.")


# ────────────────────────────────────────────────────────────────────────────
# 메인 레이아웃
# ────────────────────────────────────────────────────────────────────────────

st.title("📄 실손의료보험 초안 작성 에이전트")
st.caption(
    "삼성화재·현대해상·DB손해보험 실제 데이터를 기반으로 "
    "약관·상품설명서·사업방법서 초안을 생성하고 법령 검증을 수행합니다."
)

# 모드 토글은 모든 스텝에서 항상 렌더링해야 한다 — Streamlit은 렌더링되지 않은
# 위젯의 session_state 값을 다음 run에서 제거하므로, STEP 1에만 두면 다른
# 스텝으로 이동한 순간 워크플로우 모드가 기본값으로 되돌아간다.
st.radio(
    "약관 작성 방식",
    [WORKFLOW_MODE_GENERATE, WORKFLOW_MODE_REVISE],
    key="workflow_mode",
    horizontal=True,
    on_change=lambda: st.session_state.update(current_step=0),
    help="AI 생성: 조건을 입력하면 약관을 새로 작성합니다. "
         "직접 작성: 이미 작성된 약관 파일을 업로드하면 생성 단계 없이 법규 검증·수정만 수행합니다.",
)
st.divider()

col_form, col_result = st.columns([4, 6], gap="large")

generate_btn = False
run_all_btn  = False
revise_btn   = False
is_revise_mode = st.session_state.workflow_mode == WORKFLOW_MODE_REVISE
steps = STEPS_REVISE if is_revise_mode else STEPS_GENERATE
total_steps = len(steps)

with col_form:
    current_step = min(st.session_state.current_step, total_steps - 1)
    render_step_bar(current_step, steps)
    st.divider()

    can_proceed = True
    if current_step == 0:
        can_proceed = render_step1()
    elif current_step == 1:
        can_proceed = render_step2_revise() if is_revise_mode else render_step2()
    elif current_step == 2:
        can_proceed = render_step3(required=not is_revise_mode)
    elif current_step == 3:
        render_step4()

    _snapshot_form_fields()

    st.divider()
    col_prev, col_next = st.columns(2)

    with col_prev:
        if current_step > 0:
            st.button("← 이전", on_click=go_prev, use_container_width=True)

    with col_next:
        if current_step < total_steps - 1:
            st.button("다음 →", type="primary", on_click=go_next,
                      use_container_width=True, disabled=not can_proceed)
        elif is_revise_mode:
            can_gen = bool(st.session_state.get("user_document_text", ""))
            revise_btn = st.button(
                "🔍 검증·수정 시작",
                type="primary", use_container_width=True, disabled=not can_gen,
                help="업로드한 약관을 법규 기준으로 검증하고 위반 항목만 수정합니다 (생성 단계 없음).",
            )
            if not can_gen:
                st.warning("⚠️ STEP 2에서 약관 파일을 업로드하세요.")
        else:
            can_gen = bool(st.session_state.get("basic_coverage_items", []))

            generate_btn = st.button(
                "⚡ 기능 확인 (Upstage)",
                type="primary", use_container_width=True, disabled=not can_gen,
                help="Upstage Solar-Pro로 빠르게 기능 확인",
            )
            run_all_btn = st.button(
                "🚀 최종 실행",
                use_container_width=True, disabled=not can_gen,
                help="에이전트별 최적 모델 자동 선택",
            )
            if not can_gen:
                st.warning("⚠️ STEP 3에서 기본 보장 종목을 선택하세요.")

# ────────────────────────────────────────────────────────────────────────────
# 결과 패널
# ────────────────────────────────────────────────────────────────────────────

with col_result:
    st.subheader("검증·수정 결과" if is_revise_mode else "생성된 초안")

    if generate_btn or run_all_btn or revise_btn:
        if (not is_revise_mode
                and st.session_state.get("fetal_enrollment") == "불가"
                and st.session_state.get("join_age_min") == "태아"):
            st.error("⚠️ 태아 가입 불가 설정이지만 최소 가입나이가 태아입니다. STEP 2를 확인하세요.")
        else:
            try:
                if revise_btn:
                    with st.status("약관 검증·수정 중...", expanded=True) as status_box:
                        request = build_revise_request(model=None)
                        response = requests.post(
                            f"{BACKEND_URL}/generate/stream",
                            json=request, stream=True, timeout=300,
                        )
                        response.raise_for_status()
                        result = None
                        for event in sseclient.SSEClient(response).events():
                            if event.data == "[DONE]":
                                break
                            data = json.loads(event.data)
                            if data.get("type") == "progress":
                                node = data.get("node", "")
                                accuracy = data.get("accuracy")
                                if node == "compliance" and accuracy is not None:
                                    color = "🟢" if accuracy >= 80 else ("🟡" if accuracy >= 60 else "🔴")
                                    st.write(f"⚖️ **compliance** 완료 → {color} 준수율: **{accuracy}%**")
                                else:
                                    node_emoji = {
                                        "coordinator": "🔍", "planner": "📋", "supervisor": "🎯",
                                        "generation": "✍️", "edit": "✏️", "final_validation": "🔎",
                                        "revise": "🔧",
                                    }.get(node, "⚙️")
                                    if node:
                                        st.write(f"{node_emoji} **{node}** 완료")
                            elif data.get("type") == "result":
                                result = data
                            elif data.get("type") == "error":
                                raise Exception(data.get("message", "스트리밍 오류"))
                        if result is None:
                            raise Exception("결과를 받지 못했습니다.")
                        model_used = result.get("model_used", "Upstage Solar")
                        st.write(f"✅ 검증·수정 완료: {model_used}")
                elif generate_btn:
                    with st.status("초안 생성 중...", expanded=True) as status_box:
                        request = build_request(model=None)
                        response = requests.post(
                            f"{BACKEND_URL}/generate/stream",
                            json=request, stream=True, timeout=300,
                        )
                        response.raise_for_status()
                        result = None
                        for event in sseclient.SSEClient(response).events():
                            if event.data == "[DONE]":
                                break
                            data = json.loads(event.data)
                            if data.get("type") == "progress":
                                node = data.get("node", "")
                                accuracy = data.get("accuracy")
                                if node == "compliance" and accuracy is not None:
                                    color = "🟢" if accuracy >= 80 else ("🟡" if accuracy >= 60 else "🔴")
                                    st.write(f"⚖️ **compliance** 완료 → {color} 준수율: **{accuracy}%**")
                                else:
                                    node_emoji = {
                                        "coordinator": "🔍", "planner": "📋", "supervisor": "🎯",
                                        "generation": "✍️", "edit": "✏️", "final_validation": "🔎",
                                        "revise": "🔧",
                                    }.get(node, "⚙️")
                                    if node:
                                        st.write(f"{node_emoji} **{node}** 완료")
                            elif data.get("type") == "result":
                                result = data
                            elif data.get("type") == "error":
                                raise Exception(data.get("message", "스트리밍 오류"))
                        if result is None:
                            raise Exception("결과를 받지 못했습니다.")
                        model_used = result.get("model_used", "Upstage Solar")
                        st.write(f"✅ 생성 완료: {model_used}")
                else:
                    with st.status("최종 실행 중...", expanded=True) as status_box:
                        request = build_request(model=None)
                        response = requests.post(
                            f"{BACKEND_URL}/generate/stream",
                            json=request, stream=True, timeout=300,
                        )
                        response.raise_for_status()
                        result = None
                        for event in sseclient.SSEClient(response).events():
                            if event.data == "[DONE]":
                                break
                            data = json.loads(event.data)
                            if data.get("type") == "progress":
                                node = data.get("node", "")
                                accuracy = data.get("accuracy")
                                if node == "compliance" and accuracy is not None:
                                    color = "🟢" if accuracy >= 80 else ("🟡" if accuracy >= 60 else "🔴")
                                    st.write(f"⚖️ **compliance** 완료 → {color} 준수율: **{accuracy}%**")
                                else:
                                    node_emoji = {
                                        "coordinator": "🔍", "planner": "📋", "supervisor": "🎯",
                                        "generation": "✍️", "edit": "✏️", "final_validation": "🔎",
                                        "revise": "🔧",
                                    }.get(node, "⚙️")
                                    if node:
                                        st.write(f"{node_emoji} **{node}** 완료")
                            elif data.get("type") == "result":
                                result = data
                            elif data.get("type") == "error":
                                raise Exception(data.get("message", "스트리밍 오류"))
                        if result is None:
                            raise Exception("결과를 받지 못했습니다.")
                        model_used = result.get("model_used", "에이전트별 최적 모델")
                        st.write(f"✅ 최종 실행 완료: {model_used}")

                final_status = result.get("status", "")
                action_word = "검토" if revise_btn else "재생성"
                label = (
                    f"{'검증·수정' if revise_btn else '초안 생성'} 완료 — {result.get('iteration','?')}회 검토 통과 ({model_used})"
                    if final_status == "COMPLIANCE_PASSED"
                    else f"최대 {MAX_ITERATIONS}회 도달 — 수동 검토 필요"
                    if final_status == "MANUAL_REVIEW_REQUIRED"
                    else "오류 발생"
                )
                state = (
                    "complete" if final_status == "COMPLIANCE_PASSED" else
                    "error"    if final_status in ("MANUAL_REVIEW_REQUIRED", "ORCHESTRATOR_ERROR")
                    else "running"
                )
                status_box.update(label=label, state=state)
                st.session_state.generation_result = result
                st.session_state.generation_model  = model_used
                st.session_state.generation_show_all_docs = not revise_btn
                st.session_state.regulatory_risk_result = None

            except requests.exceptions.HTTPError as e:
                st.error(f"오류 발생: {e}")
                if e.response is not None:
                    st.json(e.response.json())
            except Exception as e:
                st.error(f"오류 발생: {e}")
                st.exception(e)

    if "generation_result" in st.session_state:
        render_result_panel(
            st.session_state.generation_result,
            st.session_state.get("generation_model", "Upstage Solar"),
            show_all_docs=st.session_state.get("generation_show_all_docs", True),
        )
    elif not (generate_btn or run_all_btn or revise_btn):
        company = st.session_state.get("insurance_company", "삼성화재")
        if is_revise_mode:
            st.info(
                f"**현재 설정:** {company}\n\n"
                "좌측에서 약관 파일을 업로드한 후 **STEP 4**에서 검증·수정을 시작하세요.\n\n"
                "AI 생성 단계 없이 법규 검증 → 수정을 반복하여 최종 약관만 출력합니다."
            )
            st.markdown("""
**검증·수정 흐름**
1. 🔍 **Coordinator** — 요청 유효성 검증
2. 📋 **Planner** — 실행 계획 수립
3. ⚖️ **Compliance** — 업로드된 약관 법규 검증
4. 🎯 **Supervisor** — 통과/재검증 결정
5. 🔧 **Revise** — 위반 항목만 수정 후 재검증 (반복)
            """)
        else:
            st.info(
                f"**현재 설정:** {company}\n\n"
                "좌측 폼을 모두 입력한 후 **STEP 4**에서 생성 버튼을 클릭하세요.\n\n"
                "약관·상품설명서·사업방법서 3개 문서가 한 번에 생성됩니다."
            )
            st.markdown("""
**생성 흐름 (LangManus 아키텍처)**
1. 🔍 **Coordinator** — 요청 유효성 검증
2. 📋 **Planner** — 실행 계획 수립
3. ✍️ **Generation** — RAG + LLM으로 초안 생성
4. ⚖️ **Compliance** — 법규 준수 검증 (5개 룰)
5. 🎯 **Supervisor** — 재생성 or 편집 결정
6. ✏️ **Edit** — 최종 편집 및 다중 문서 생성
            """)
        with st.expander("📊 3사 데이터 비교"):
            st.markdown("""
| 항목 | 삼성화재 | 현대해상 | DB손해보험 |
|------|---------|---------|-----------|
| 가입나이 | 0세~65세 | 태아~60세 | 5세~99세 |
| MRI 한도 | 300만원 | 300만원 | 200만원 |
| 도수치료 | 350만원 | 350만원 | 350만원 |
| 보험기간 | 1년 만기 | 1년 만기 | 1년 만기 |
| 갱신주기 | 1년 | 1년 | 1년 |
| 재가입주기 | 5년 | 5년 | 5년 |
            """)
